"""
Resume Parser Module
Extracts text from PDF and DOCX resumes for analysis
"""

import pdfplumber
import os
import re
from docx import Document

def extract_resume_text(file_path):
    """
    Extract text from resume (PDF or DOCX)
    
    Args:
        file_path: Path to PDF or DOCX file
    
    Returns:
        Extracted text from resume
    """
    try:
        if not os.path.exists(file_path):
            print(f"❌ File not found: {file_path}")
            return ""
        
        # Determine file type by extension
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return extract_pdf_text(file_path)
        elif file_ext in ['.docx', '.doc']:
            return extract_docx_text(file_path)
        else:
            print(f"❌ Unsupported file format: {file_ext}. Only PDF and DOCX are supported.")
            return ""
    
    except Exception as e:
        print(f"❌ Error parsing resume: {e}")
        return ""

def extract_pdf_text(file_path):
    """Extract text from PDF file"""
    try:
        text = ""
        
        with pdfplumber.open(file_path) as pdf:
            # Extract text from all pages
            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text
                    text += "\n--- PAGE BREAK ---\n"
        
        if not text.strip():
            print(f"⚠️  No text extracted from {file_path}")
            return ""
        
        print(f"✅ Resume parsed: {file_path} ({len(text)} characters)")
        return text
    
    except Exception as e:
        print(f"❌ Error parsing PDF: {e}")
        return ""

def extract_docx_text(file_path):
    """Extract text from DOCX file"""
    try:
        doc = Document(file_path)
        text = ""
        
        # Extract text from all paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        if not text.strip():
            print(f"⚠️  No text extracted from {file_path}")
            return ""
        
        print(f"✅ Resume parsed: {file_path} ({len(text)} characters)")
        return text
    
    except Exception as e:
        print(f"❌ Error parsing DOCX: {e}")
        return ""

def extract_candidate_info(resume_text):
    """
    Extract basic candidate information from resume
    
    Args:
        resume_text: Full text from resume
    
    Returns:
        Dictionary with candidate info
    """
    try:
        info = {
            'emails': extract_emails(resume_text),
            'phones': extract_phones(resume_text),
            'linkedin': extract_linkedin(resume_text),
            'skills': extract_skills(resume_text),
            'experience_years': extract_experience_years(resume_text)
        }
        
        return info
    
    except Exception as e:
        print(f"❌ Error extracting candidate info: {e}")
        return {}

def extract_emails(text):
    """Extract email addresses from text"""
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    return re.findall(email_pattern, text)

def extract_phones(text):
    """Extract phone numbers from text"""
    phone_pattern = r'\+?[\d\s\-\(\)]{10,}'
    return re.findall(phone_pattern, text)

def extract_linkedin(text):
    """Extract LinkedIn profile URL"""
    linkedin_pattern = r'linkedin\.com/in/[\w\-]+'
    matches = re.findall(linkedin_pattern, text)
    return matches[0] if matches else None

def extract_skills(text):
    """Extract technical skills mentioned in resume"""
    common_skills = [
        'Python', 'Java', 'JavaScript', 'C++', 'C#', 'PHP', 'Ruby', 'Go', 'SQL',
        'React', 'Angular', 'Vue', 'Node.js', 'Django', 'Flask', 'Spring',
        'AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes', 'Git',
        'Machine Learning', 'Data Analysis', 'AI', 'Deep Learning',
        'Excel', 'Tableau', 'Power BI', 'Salesforce',
        'Project Management', 'Agile', 'Scrum', 'Communication'
    ]
    
    found_skills = []
    for skill in common_skills:
        if skill.lower() in text.lower():
            found_skills.append(skill)
    
    return found_skills

def extract_experience_years(text):
    """Estimate years of experience from resume"""
    # Look for patterns like "5 years of experience" or "2018-2023"
    year_pattern = r'(\d{1,2})\s+years?\s+(?:of\s+)?experience'
    matches = re.findall(year_pattern, text.lower())
    
    if matches:
        try:
            return int(matches[0])
        except:
            pass
    
    # Try to count years from date ranges
    date_pattern = r'(\d{4})\s*-\s*(?:Present|\d{4})'
    dates = re.findall(date_pattern, text)
    if dates:
        try:
            return len(dates)
        except:
            pass
    
    return 0

def get_resume_filename(file_path):
    """Extract just the filename from path"""
    return os.path.basename(file_path)

def extract_candidate_name_from_filename(file_path):
    """Extract and clean candidate name from filename"""
    filename = os.path.basename(file_path)
    # Remove all file extensions (.pdf, .docx, .doc)
    name = filename.rsplit('.', 1)[0]
    # Remove version numbers like (1), (2), etc.
    name = re.sub(r'\s*\(\d+\)\s*', ' ', name)
    # Replace underscores and hyphens with spaces
    name = name.replace('_', ' ').replace('-', ' ')
    # Remove common resume keywords like 'resume', 'cv', 'cv_', etc.
    name = re.sub(r'\s*(resume|cv|cv_|application|candidate)\s*', ' ', name, flags=re.IGNORECASE)
    # Clean up multiple spaces
    name = ' '.join(name.split())
    # Capitalize properly
    name = name.title().strip()
    return name if name else "Candidate"

def extract_candidate_name_from_text(resume_text):
    """Try to extract name from resume text itself"""
    lines = resume_text.split('\n')
    
    # Common institution/company names to skip
    skip_patterns = [
        'email', 'phone', 'address', 'experience', 'education', 'objective',
        'university', 'college', 'institute', 'school', 'academy', 'institute of',
        'sibm', 'iit', 'mit', 'stanford', 'harvard', 'delhi', 'mumbai', 'pune',
        'sistec', 'nvcc', 'bsnl', 'accenture', 'cognizant', 'infosys', 'tcs',
        'linkedin', 'github', 'portfolio', 'website', 'summary', 'profile', 'about',
        'skills', 'projects', 'internship', 'certification', 'languages', 'hobbies'
    ]
    
    # First, look for explicit "Name:" patterns
    for line in lines[:20]:
        if 'name:' in line.lower() or 'name -' in line.lower():
            # Extract text after "Name:"
            match = re.search(r'(?:name\s*[:|-]?\s*)([A-Za-z\s\.]+)', line, re.IGNORECASE)
            if match:
                name = match.group(1).strip().title()
                if len(name.split()) >= 2 and len(name) < 60:
                    return name
    
    # Check first 15 lines for candidate name
    for line in lines[:15]:
        line = line.strip()
        
        # Skip empty lines and known non-name lines
        if not line or len(line) > 100:
            continue
        
        # Skip if line contains skip keywords
        if any(keyword in line.lower() for keyword in skip_patterns):
            continue
        
        # Likely a name if it's 2-4 words, all capitalized or title case
        words = line.split()
        
        if 2 <= len(words) <= 4:
            # Check if has typical name patterns
            # Names usually start with capital letters
            if all(word[0].isupper() for word in words if word):
                # Filter out lines with numbers, special chars (except hyphen, apostrophe)
                if not re.search(r'[0-9@#$%^&*()+=\[\]{}|;:,<>?/\\~`]', line):
                    return line.title()
    
    return None

if __name__ == "__main__":
    # Test the parser
    print("Resume Parser Test\n")
    
    test_resumes = [
        "resumes/sample_resume.pdf",
    ]
    
    for resume in test_resumes:
        if os.path.exists(resume):
            text = extract_resume_text(resume)
            if text:
                info = extract_candidate_info(text)
                print(f"\nCandidate Info from {resume}:")
                print(f"  Emails: {info.get('emails')}")
                print(f"  Phones: {info.get('phones')}")
                print(f"  LinkedIn: {info.get('linkedin')}")
                print(f"  Skills: {info.get('skills')}")
                print(f"  Experience (years): {info.get('experience_years')}")